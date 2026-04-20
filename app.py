import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# ================================
# 🌙 DARK MODE TOGGLE
# ================================
mode = st.sidebar.toggle("🌙 Dark Mode")

if mode:
    bg_color = "#1e1e1e"
    text_color = "white"
else:
    bg_color = "#f5f7fa"
    text_color = "black"

# ================================
# 🎨 CUSTOM CSS
# ================================
st.markdown(f"""
<style>
.stApp {{
    background-color: {bg_color};
    color: {text_color};
}}

h1 {{
    text-align: center;
}}

.stButton>button {{
    background-color: #3498db;
    color: white;
    border-radius: 10px;
}}

.stDownloadButton>button {{
    background-color: #27ae60;
    color: white;
    border-radius: 10px;
}}

[data-testid="metric-container"] {{
    background-color: #ffffff;
    border-radius: 12px;
    padding: 15px;
    box-shadow: 0px 4px 12px rgba(0,0,0,0.15);
}}
</style>
""", unsafe_allow_html=True)

# ================================
# 🖼️ LOGO + TITLE
# ================================
st.markdown("<h1>📊 Publication Summary Generator</h1>", unsafe_allow_html=True)

st.sidebar.title("📌 Navigation")
page = st.sidebar.radio("Go to", ["🏠 Home", "📊 Dashboard"])

# ================================
# 🏠 HOME PAGE
# ================================
if page == "🏠 Home":
    st.write("Upload your publication data to generate reports.")

    uploaded_file = st.file_uploader("Upload publications.xlsx", type=["xlsx"])

    if uploaded_file:
        df = pd.read_excel(uploaded_file)
        df.columns = ["Faculty Name", "Year", "Title", "Venue"]
        df["Year"] = pd.to_numeric(df["Year"], errors="coerce")

        def classify(v):
            v = str(v).lower()
            if "journal" in v:
                return "Journal"
            elif "conference" in v or "conf" in v:
                return "Conference"
            return "Other"

        df["Type"] = df["Venue"].apply(classify)

        st.session_state["data"] = df
        st.success("✅ File uploaded successfully! Go to Dashboard")

# ================================
# 📊 DASHBOARD
# ================================
if page == "📊 Dashboard":

    if "data" not in st.session_state:
        st.warning("⚠️ Please upload file first from Home page")
    else:
        df = st.session_state["data"]

        st.subheader("📅 Select Year Range")

        min_year = int(df["Year"].min())
        max_year = int(df["Year"].max())

        start_year, end_year = st.slider(
            "Select range",
            min_year,
            max_year,
            (min_year, max_year)
        )

        df = df[(df["Year"] >= start_year) & (df["Year"] <= end_year)]

        # ================================
        # 📈 METRICS
        # ================================
        col1, col2, col3 = st.columns(3)
        col1.metric("Total", len(df))
        col2.metric("Journals", len(df[df["Type"] == "Journal"]))
        col3.metric("Conferences", len(df[df["Type"] == "Conference"]))

        # ================================
        # 📊 GRAPH
        # ================================
        st.subheader("📊 Publications per Year")
        chart_data = df.groupby("Year").size()
        st.bar_chart(chart_data)

        # ================================
        # 📄 TABLES
        # ================================
        st.subheader("📄 Journal Publications")
        st.dataframe(df[df["Type"] == "Journal"])

        st.subheader("📄 Conference Publications")
        st.dataframe(df[df["Type"] == "Conference"])

        # ================================
        # 📥 DOWNLOAD EXCEL
        # ================================
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)

        st.download_button("📥 Download Excel", excel_buffer.getvalue(), "report.xlsx")

        # ================================
        # 📥 DOWNLOAD WORD
        # ================================
        doc = Document()
        doc.add_heading("Faculty Publication Report", 0)

        for name in df["Faculty Name"].unique():
            doc.add_heading(name, level=1)
            for _, row in df[df["Faculty Name"] == name].iterrows():
                doc.add_paragraph(f"{int(row['Year'])} - {row['Title']} ({row['Type']})")

        word_buffer = BytesIO()
        doc.save(word_buffer)

        st.download_button("📥 Download Word", word_buffer.getvalue(), "report.docx")