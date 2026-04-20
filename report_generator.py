import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.chart import BarChart, Reference

# ================================
# LOAD DATA
# ================================
df = pd.read_excel("publications.xlsx")
df.columns = ["Faculty Name", "Year", "Title", "Venue"]

df["Year"] = pd.to_numeric(df["Year"], errors="coerce")

# ================================
# CLASSIFY
# ================================
def classify(venue):
    venue = str(venue).lower()
    if "journal" in venue:
        return "Journal"
    elif "conference" in venue or "conf" in venue:
        return "Conference"
    else:
        return "Other"

df["Type"] = df["Venue"].apply(classify)

# ================================
# FILTER
# ================================
start_year = 2020
end_year = 2025

df = df[(df["Year"] >= start_year) & (df["Year"] <= end_year)]

journal_df = df[df["Type"] == "Journal"]
conference_df = df[df["Type"] == "Conference"]

# ================================
# SUMMARY
# ================================
summary = df.groupby("Year").size().reset_index(name="Total Publications")

# ================================
# SAVE EXCEL
# ================================
with pd.ExcelWriter("final_report.xlsx", engine='openpyxl') as writer:
    df.to_excel(writer, sheet_name="All Data", index=False)
    journal_df.to_excel(writer, sheet_name="Journals", index=False)
    conference_df.to_excel(writer, sheet_name="Conferences", index=False)
    summary.to_excel(writer, sheet_name="Summary", index=False)

# ================================
# FORMAT EXCEL
# ================================
wb = load_workbook("final_report.xlsx")

header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
header_font = Font(color="FFFFFF", bold=True)

for sheet in wb.sheetnames:
    ws = wb[sheet]

    # Header styling
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # Auto width
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter

        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))

        ws.column_dimensions[col_letter].width = max_length + 3

# ================================
# ADD CHART
# ================================
ws = wb["Summary"]

chart = BarChart()
chart.title = "Publications Per Year"

data = Reference(ws, min_col=2, min_row=1, max_row=len(summary)+1)
cats = Reference(ws, min_col=1, min_row=2, max_row=len(summary)+1)

chart.add_data(data, titles_from_data=True)
chart.set_categories(cats)

ws.add_chart(chart, "E5")

wb.save("final_report.xlsx")

print("✅ Premium Excel report created!")

# ================================
# WORD REPORT (PREMIUM)
# ================================
doc = Document()

doc.add_heading('Faculty Publication Report', 0)
doc.add_paragraph(f'Duration: {start_year} to {end_year}')

doc.add_heading('Journal Publications', level=1)

for name in journal_df["Faculty Name"].unique():
    doc.add_heading(name, level=2)
    for _, row in journal_df[journal_df["Faculty Name"] == name].iterrows():
        doc.add_paragraph(f"{int(row['Year'])} - {row['Title']}")

doc.add_heading('Conference Publications', level=1)

for name in conference_df["Faculty Name"].unique():
    doc.add_heading(name, level=2)
    for _, row in conference_df[conference_df["Faculty Name"] == name].iterrows():
        doc.add_paragraph(f"{int(row['Year'])} - {row['Title']}")

doc.save("final_report.docx")

print("✅ Premium Word report created!")